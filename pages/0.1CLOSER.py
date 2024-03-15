import streamlit as st
import pandas as pd
import docx
import fitz  # PyMuPDF
from langchain.llms import openai
from langchain.chains import qa_generation  

from langchain_community.chat_models import ChatOpenAI
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain


# Initialize Langchain with OpenAI
with st.sidebar:
     openai_api_key = st.text_input("OpenAI API Key", key="chatbot_api_key", type="password")
"[Get an OpenAI API key](https://platform.openai.com/account/api-keys)"
    
def read_pdf(file):
    with fitz.open(stream=file, filetype="pdf") as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

def read_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def upload_business_process_document():
    uploaded_file = st.file_uploader("Upload your business process document", type=['txt', 'pdf', 'docx'])
    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            content = read_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content = read_docx(uploaded_file)
        else:  # Assuming text file
            content = uploaded_file.getvalue().decode("utf-8")
        return content
    return None

def upload_detailed_steps_documents():
    uploaded_files = st.file_uploader("Upload detailed steps documents for each activity", accept_multiple_files=True, type=['txt', 'pdf', 'docx'])
    detailed_docs = {}
    for uploaded_file in uploaded_files:
        if uploaded_file.type == "application/pdf":
            content = read_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content = read_docx(uploaded_file)
        else:  # Assuming text file
            content = uploaded_file.getvalue().decode("utf-8")
        detailed_docs[uploaded_file.name] = content
    return detailed_docs

def process_documents(business_process_doc, detailed_steps_docs):
    
    prompt = ChatPromptTemplate.from_messages(
    [("system", "Generate a test cases using the \n\n{context}")]
)
#    
    llm = ChatOpenAI(model_name="gpt-3.5-turbo",openai_api_key=openai_api_key)
    chain = create_stuff_documents_chain(llm, prompt)

    docs = [Document(page_content= business_process_doc),
            Document(page_content= detailed_steps_docs)]

    chain.invoke({"context": docs})

#    combined_text = business_process_doc + "\n\n" + "\n\n".join(detailed_steps_docs.values())
#    prompt = f"Based on the following business process and detailed steps, generate comprehensive test cases:\n\n{combined_text}\n\nTest cases:"
#    response = generation_chain.run(prompt=prompt, max_tokens=1024)  # Adjust max_tokens as needed
#    test_cases = response.split('\n')  # Assuming each test case is separated by a newline
#    return test_cases

def display_test_cases(test_cases):
    st.subheader("Generated Test Cases")
    for i, test_case in enumerate(test_cases, start=1):
        st.text(f"{i}. {test_case}")

    # Optionally, allow users to download the test cases as a file
    if st.button('Download Test Cases'):
        test_cases_df = pd.DataFrame(test_cases, columns=["Test Cases"])
        csv = test_cases_df.to_csv(index=False)
        st.download_button(label="Download test cases as CSV", data=csv, file_name='test_cases.csv', mime='text/csv')

def main():

    st.title("Test Case Generator for Business Processes")

    business_process_doc = upload_business_process_document()
    detailed_steps_docs = upload_detailed_steps_documents()

    if st.button("Generate Test Cases") and business_process_doc and detailed_steps_docs:
        test_cases = process_documents(business_process_doc, detailed_steps_docs)
#        display_test_cases(test_cases)

if __name__ == "__main__":
    main()
